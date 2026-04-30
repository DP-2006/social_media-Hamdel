
# import os
# import logging
# from django.conf import settings
# from django.core.cache import cache

# logger = logging.getLogger(__name__)

# try:
#     import PyPDF2
# except ImportError:
#     PyPDF2 = None
#     logger.warning("PyPDF2 not installed. PDF files will not be analyzed.")

# try:
#     import docx
# except ImportError:
#     docx = None
#     logger.warning("python-docx not installed. DOCX files will not be analyzed.")

# try:
#     from PIL import Image
# except ImportError:
#     Image = None
#     logger.warning("Pillow not installed. Images will not be analyzed.")

# try:
#     import openpyxl
# except ImportError:
#     openpyxl = None
#     logger.warning("openpyxl not installed. Excel files will not be analyzed.")


# # class FileContentAnalyzer:
    
# #     @staticmethod
# #     def extract_text_from_pdf(file_path):
# #         try:
# #             text = ""
# #             with open(file_path, 'rb') as file:
# #                 pdf_reader = PyPDF2.PdfReader(file)
# #                 for page in pdf_reader.pages[:10]: 
# #                     text += page.extract_text() or ""
# #             return text[:2000]
# #         except Exception as e:
# #             logger.error(f"Error reading PDF: {e}")
# #             return None
    
# #     @staticmethod
# #     def extract_text_from_docx(file_path):
# #         try:
# #             doc = docx.Document(file_path)
# #             text = "\n".join([paragraph.text for paragraph in doc.paragraphs[:50]])
# #             return text[:2000]
# #         except Exception as e:
# #             logger.error(f"Error reading DOCX: {e}")
# #             return None
    
# #     @staticmethod
# #     def extract_text_from_txt(file_path):
# #         try:
# #             with open(file_path, 'r', encoding='utf-8') as file:
# #                 text = file.read()
# #             return text[:2000]
# #         except UnicodeDecodeError:
# #             try:
# #                 with open(file_path, 'r', encoding='latin-1') as file:
# #                     text = file.read()
# #                 return text[:2000]
# #             except Exception as e:
# #                 logger.error(f"Error reading TXT: {e}")
# #                 return None
# #         except Exception as e:
# #             logger.error(f"Error reading TXT: {e}")
# #             return None
    
# #     @staticmethod
# #     def extract_metadata_from_image(file_path):
# #         try:
# #             with Image.open(file_path) as img:
# #                 metadata = {
# #                     "format": img.format,
# #                     "size": img.size,
# #                     "mode": img.mode,
# #                     "width": img.width,
# #                     "height": img.height
# #                 }
# #                 return metadata
# #         except Exception as e:
# #             logger.error(f"Error reading image: {e}")
# #             return None
    
# #     @staticmethod
# #     def get_file_summary(file_obj):
# #         file_path = file_obj.path
# #         file_extension = os.path.splitext(file_obj.name)[1].lower()
        
# #         result = {
# #             "file_name": file_obj.name,
# #             "file_size": file_obj.size,
# #             "file_type": file_extension,
# #             "content_summary": None,
# #             "key_topics": [],
# #             "metadata": {}
# #         }
        
# #         if file_extension == '.pdf':
# #             content = FileContentAnalyzer.extract_text_from_pdf(file_path)
# #             if content:
# #                 result["content_summary"] = content[:500]
# #                 result["key_topics"] = FileContentAnalyzer._extract_key_topics(content)
                
# #         elif file_extension == '.docx':
# #             content = FileContentAnalyzer.extract_text_from_docx(file_path)
# #             if content:
# #                 result["content_summary"] = content[:500]
# #                 result["key_topics"] = FileContentAnalyzer._extract_key_topics(content)
                
# #         elif file_extension == '.txt':
# #             content = FileContentAnalyzer.extract_text_from_txt(file_path)
# #             if content:
# #                 result["content_summary"] = content[:500]
# #                 result["key_topics"] = FileContentAnalyzer._extract_key_topics(content)
                
# #         elif file_extension in ['.xlsx', '.xls']:
# #             result["content_summary"] = "فایل اکسل - داده‌های جدولی"
                
# #         elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
# #             metadata = FileContentAnalyzer.extract_metadata_from_image(file_path)
# #             if metadata:
# #                 result["metadata"] = metadata
# #                 result["content_summary"] = f"تصویر با فرمت {metadata.get('format')} و ابعاد {metadata.get('width')}x{metadata.get('height')}"
        
# #         return result
    
# #     @staticmethod
# #     def _extract_key_topics(text, max_topics=5):
# #         words = text.lower().split()
# #         word_count = {}
# #         stop_words = {'و', 'از', 'به', 'را', 'با', 'برای', 'در', 'این', 'آن', 'که', 'نیست', 'است'}
        
# #         for word in words:
# #             if len(word) > 3 and word not in stop_words:
# #                 word_count[word] = word_count.get(word, 0) + 1
        
# #         sorted_words = sorted(word_count.items(), key=lambda x: x[1], reverse=True)
# #         return [word for word, count in sorted_words[:max_topics]]


# # class UserFilesAnalyzer:
# #     """تحلیل تمام فایل‌های یک کاربر"""
    
# #     def __init__(self, user):
# #         self.user = user
# #         self.file_analyzer = FileContentAnalyzer()
    
# #     # def analyze_all_files(self):
# #     #     """تحلیل تمام فایل‌های کاربر از پست‌ها"""
# #     #     try:
# #     #         from apps.posts.models import Post
            
# #     #         posts_with_files = Post.objects.filter(
# #     #             user=self.user,
# #     #             file__isnull=False
# #     #         ).exclude(file='')
            
# #     #         files_analysis = []
# #     #         all_content = []
# #     #         file_types = {}
# #     #         total_size = 0
            
# #     #         for post in posts_with_files:
# #     #             try:
# #     #                 if post.file and hasattr(post.file, 'path') and os.path.exists(post.file.path):
# #     #                     file_info = self.file_analyzer.get_file_summary(post.file)
                        
# #     #                     files_analysis.append({
# #     #                         "post_id": str(post.id),
# #     #                         "post_content": post.content[:100] if post.content else "",
# #     #                         "file_info": file_info,
# #     #                         "created_at": post.created_at.isoformat() if post.created_at else None
# #     #                     })
                        
# #     #                     ext = file_info["file_type"]
# #     #                     file_types[ext] = file_types.get(ext, 0) + 1
# #     #                     total_size += file_info["file_size"]
                        
# #     #                     if file_info.get("content_summary"):
# #     #                         all_content.append(file_info["content_summary"])
                            
# #     #             except Exception as e:
# #     #                 logger.error(f"Error analyzing file for post {post.id}: {e}")
# #     #                 files_analysis.append({
# #     #                     "post_id": str(post.id),
# #     #                     "error": f"خطا در تحلیل: {str(e)}"
# #     #                 })
            
# #     #         return {
# #     #             "total_files": posts_with_files.count(),
# #     #             "total_size_bytes": total_size,
# #     #             "file_types": file_types,
# #     #             "files_analysis": files_analysis[:10],
# #     #             "combined_content": " ".join(all_content)[:2000]
# #     #         }
            
# #     #     except Exception as e:
# #     #         logger.error(f"Error in analyze_all_files: {e}")
# #     #         return {
# #     #             "total_files": 0,
# #     #             "total_size_bytes": 0,
# #     #             "file_types": {},
# #     #             "files_analysis": [],
# #     #             "combined_content": "",
# #     #             "error": str(e)
# #     #         }


# # # apps/ml/services/file_analyzer.py
# # # متد analyze_all_files رو اصلاح کن:

# # def analyze_all_files(self):
# #     """تحلیل تمام فایل‌های کاربر از پست‌ها"""
# #     try:
# #         from apps.posts.models import Post
        
# #         posts_with_files = Post.objects.filter(
# #             user=self.user,
# #             file__isnull=False
# #         ).exclude(file='')
        
# #         files_analysis = []
# #         all_content = []
# #         file_types = {}
# #         total_size = 0
        
# #         for post in posts_with_files:
# #             try:
# #                 if post.file and hasattr(post.file, 'path'):
# #                     # چک کن فایل واقعاً وجود داره
# #                     if os.path.exists(post.file.path):
# #                         file_info = self.file_analyzer.get_file_summary(post.file)
                        
# #                         files_analysis.append({
# #                             "post_id": str(post.id),
# #                             "post_content": post.content[:100] if post.content else "",
# #                             "file_info": file_info,
# #                             "created_at": post.created_at.isoformat() if post.created_at else None
# #                         })
                        
# #                         ext = file_info["file_type"]
# #                         file_types[ext] = file_types.get(ext, 0) + 1
# #                         total_size += file_info["file_size"]
                        
# #                         if file_info.get("content_summary"):
# #                             all_content.append(file_info["content_summary"])
# #                     else:
# #                         # فایل وجود نداره
# #                         files_analysis.append({
# #                             "post_id": str(post.id),
# #                             "error": f"فایل {post.file.name} وجود ندارد"
# #                         })
                        
# #             except Exception as e:
# #                 logger.error(f"Error analyzing file for post {post.id}: {e}")
# #                 files_analysis.append({
# #                     "post_id": str(post.id),
# #                     "error": f"خطا در تحلیل: {str(e)}"
# #                 })
        
# #         return {
# #             "total_files": posts_with_files.count(),
# #             "total_size_bytes": total_size,
# #             "file_types": file_types,
# #             "files_analysis": files_analysis[:10],
# #             "combined_content": " ".join(all_content)[:2000]
# #         }
        
# #     except Exception as e:
# #         logger.error(f"Error in analyze_all_files: {e}")
# #         return {
# #             "total_files": 0,
# #             "total_size_bytes": 0,
# #             "file_types": {},
# #             "files_analysis": [],
# #             "combined_content": "",
# #             "error": str(e)
# #         }





# # class FileContentAnalyzer:
    
# #     @staticmethod
# #     def extract_text_from_pdf(file_path):
# #         if PyPDF2 is None:
# #             return None
# #         try:
# #             text = ""
# #             with open(file_path, 'rb') as file:
# #                 pdf_reader = PyPDF2.PdfReader(file)
# #                 for page in pdf_reader.pages[:10]: 
# #                     text += page.extract_text() or ""
# #             return text[:2000]
# #         except Exception as e:
# #             logger.error(f"Error reading PDF: {e}")
# #             return None
    
# #     @staticmethod
# #     def extract_text_from_docx(file_path):
# #         if docx is None:
# #             return None
# #         try:
# #             doc = docx.Document(file_path)
# #             text = "\n".join([paragraph.text for paragraph in doc.paragraphs[:50]])
# #             return text[:2000]
# #         except Exception as e:
# #             logger.error(f"Error reading DOCX: {e}")
# #             return None
    
# #     @staticmethod
# #     def extract_metadata_from_image(file_path):
# #         if Image is None:
# #             return None
# #         try:
# #             with Image.open(file_path) as img:
# #                 metadata = {
# #                     "format": img.format,
# #                     "size": img.size,
# #                     "mode": img.mode,
# #                     "width": img.width,
# #                     "height": img.height
# #                 }
# #                 return metadata
# #         except Exception as e:
# #             logger.error(f"Error reading image: {e}")
# #             return None



















# # apps/ml/services/file_analyzer.py

# import os
# import logging
# from django.conf import settings

# logger = logging.getLogger(__name__)

# # ایمپورت‌های اختیاری
# try:
#     import PyPDF2
# except ImportError:
#     PyPDF2 = None
#     logger.warning("PyPDF2 not installed")

# try:
#     import docx
# except ImportError:
#     docx = None
#     logger.warning("python-docx not installed")

# try:
#     from PIL import Image
# except ImportError:
#     Image = None
#     logger.warning("Pillow not installed")


# class FileContentAnalyzer:
#     """تحلیل محتوای فایل‌ها"""
    
#     @staticmethod
#     def extract_text_from_pdf(file_path):
#         if PyPDF2 is None:
#             return None
#         try:
#             text = ""
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page in pdf_reader.pages[:10]:
#                     text += page.extract_text() or ""
#             return text[:2000]
#         except Exception as e:
#             logger.error(f"Error reading PDF: {e}")
#             return None
    
#     @staticmethod
#     def extract_text_from_docx(file_path):
#         if docx is None:
#             return None
#         try:
#             doc = docx.Document(file_path)
#             text = "\n".join([paragraph.text for paragraph in doc.paragraphs[:50]])
#             return text[:2000]
#         except Exception as e:
#             logger.error(f"Error reading DOCX: {e}")
#             return None
    
#     @staticmethod
#     def extract_text_from_txt(file_path):
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 text = file.read()
#             return text[:2000]
#         except UnicodeDecodeError:
#             try:
#                 with open(file_path, 'r', encoding='latin-1') as file:
#                     text = file.read()
#                 return text[:2000]
#             except Exception as e:
#                 logger.error(f"Error reading TXT: {e}")
#                 return None
#         except Exception as e:
#             logger.error(f"Error reading TXT: {e}")
#             return None
    
#     @staticmethod
#     def extract_metadata_from_image(file_path):
#         if Image is None:
#             return None
#         try:
#             with Image.open(file_path) as img:
#                 metadata = {
#                     "format": img.format,
#                     "size": img.size,
#                     "mode": img.mode,
#                     "width": img.width,
#                     "height": img.height
#                 }
#                 return metadata
#         except Exception as e:
#             logger.error(f"Error reading image: {e}")
#             return None
    
#     @staticmethod
#     def get_file_summary(file_obj):
#         file_path = file_obj.path
#         file_extension = os.path.splitext(file_obj.name)[1].lower()
        
#         result = {
#             "file_name": file_obj.name,
#             "file_size": file_obj.size,
#             "file_type": file_extension,
#             "content_summary": None,
#             "key_topics": [],
#             "metadata": {}
#         }
        
#         if file_extension == '.pdf':
#             content = FileContentAnalyzer.extract_text_from_pdf(file_path)
#             if content:
#                 result["content_summary"] = content[:500]
                
#         elif file_extension == '.docx':
#             content = FileContentAnalyzer.extract_text_from_docx(file_path)
#             if content:
#                 result["content_summary"] = content[:500]
                
#         elif file_extension == '.txt':
#             content = FileContentAnalyzer.extract_text_from_txt(file_path)
#             if content:
#                 result["content_summary"] = content[:500]
                
#         elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
#             metadata = FileContentAnalyzer.extract_metadata_from_image(file_path)
#             if metadata:
#                 result["metadata"] = metadata
#                 result["content_summary"] = f"تصویر با فرمت {metadata.get('format')}"
        
#         return result


# class UserFilesAnalyzer:
#     """تحلیل تمام فایل‌های یک کاربر"""
    
#     def __init__(self, user):
#         self.user = user
#         self.file_analyzer = FileContentAnalyzer()
    
#     # def analyze_all_files(self):
#     #     """تحلیل تمام فایل‌های کاربر از پست‌ها"""
#     #     try:
#     #         from apps.posts.models import Post
            
#     #         posts_with_files = Post.objects.filter(
#     #             user=self.user,
#     #             file__isnull=False
#     #         ).exclude(file='')
            
#     #         files_analysis = []
#     #         all_content = []
#     #         file_types = {}
#     #         total_size = 0
            
#     #         for post in posts_with_files:
#     #             try:
#     #                 if post.file and hasattr(post.file, 'path') and os.path.exists(post.file.path):
#     #                     file_info = self.file_analyzer.get_file_summary(post.file)
                        
#     #                     files_analysis.append({
#     #                         "post_id": str(post.id),
#     #                         "post_content": post.content[:100] if post.content else "",
#     #                         "file_info": file_info,
#     #                         "created_at": post.created_at.isoformat() if post.created_at else None
#     #                     })
                        
#     #                     ext = file_info["file_type"]
#     #                     file_types[ext] = file_types.get(ext, 0) + 1
#     #                     total_size += file_info["file_size"]
                        
#     #                     if file_info.get("content_summary"):
#     #                         all_content.append(file_info["content_summary"])
                            
#     #             except Exception as e:
#     #                 logger.error(f"Error analyzing file for post {post.id}: {e}")
#     #                 files_analysis.append({
#     #                     "post_id": str(post.id),
#     #                     "error": f"خطا در تحلیل: {str(e)}"
#     #                 })
            
#     #         return {
#     #             "total_files": posts_with_files.count(),
#     #             "total_size_bytes": total_size,
#     #             "file_types": file_types,
#     #             "files_analysis": files_analysis[:10],
#     #             "combined_content": " ".join(all_content)[:2000]
#     #         }
            
#     #     except Exception as e:
#     #         logger.error(f"Error in analyze_all_files: {e}")
#     #         return {
#     #             "total_files": 0,
#     #             "total_size_bytes": 0,
#     #             "file_types": {},
#     #             "files_analysis": [],
#     #             "combined_content": "",
#     #             "error": str(e)
#     #         }
#     def analyze_all_files(self):
#     try:
#         from apps.posts.models import Post, PostMedia
        
#         # پست‌های با فایل تکی
#         posts_with_single_file = Post.objects.filter(
#             user=self.user,
#             file__isnull=False
#         ).exclude(file='')
        
#         # فایل‌های چندتایی از PostMedia
#         multiple_media = PostMedia.objects.filter(
#             post__user=self.user
#         ).select_related('post')
        
#         files_analysis = []
#         all_content = []
#         file_types = {}
#         total_size = 0
        
        
#         for post in posts_with_single_file:
#             try:
#                 if post.file and hasattr(post.file, 'path') and os.path.exists(post.file.path):
#                     file_info = self.file_analyzer.get_file_summary(post.file)
                    
#                     files_analysis.append({
#                         "post_id": str(post.id),
#                         "post_content": post.content[:100] if post.content else "",
#                         "file_info": file_info,
#                         "file_source": "single_file",
#                         "created_at": post.created_at.isoformat() if post.created_at else None
#                     })
                    
#                     ext = file_info["file_type"]
#                     file_types[ext] = file_types.get(ext, 0) + 1
#                     total_size += file_info["file_size"]
                    
#                     if file_info.get("content_summary"):
#                         all_content.append(file_info["content_summary"])
                        
#             except Exception as e:
#                 logger.error(f"Error analyzing single file for post {post.id}: {e}")
        
        
#         for media in multiple_media:
#             try:
#                 if media.file and hasattr(media.file, 'path') and os.path.exists(media.file.path):
#                     file_info = self.file_analyzer.get_file_summary(media.file)
                    
#                     post = media.post
                    
#                     files_analysis.append({
#                         "post_id": str(post.id),
#                         "post_content": post.content[:100] if post.content else "",
#                         "media_id": str(media.id),
#                         "file_info": file_info,
#                         "file_source": "post_media",
#                         "order": media.order,
#                         "created_at": post.created_at.isoformat() if post.created_at else None
#                     })
                    
#                     ext = file_info["file_type"]
#                     file_types[ext] = file_types.get(ext, 0) + 1
#                     total_size += file_info["file_size"]
                    
#                     if file_info.get("content_summary"):
#                         all_content.append(file_info["content_summary"])
                        
#             except Exception as e:
#                 logger.error(f"Error analyzing media file {media.id}: {e}")
        
#         total_files = posts_with_single_file.count() + multiple_media.count()
        
#         return {
#             "total_files": total_files,
#             "total_size_bytes": total_size,
#             "file_types": file_types,
#             "files_analysis": files_analysis[:20],
#             "combined_content": " ".join(all_content)[:3000]
#         }
        
#     except Exception as e:
#         logger.error(f"Error in analyze_all_files: {e}")
#         return {
#             "total_files": 0,
#             "total_size_bytes": 0,
#             "file_types": {},
#             "files_analysis": [],
#             "combined_content": "",
#             "error": str(e)
#         }








# apps/ml/services/file_analyzer.py
import os
import logging
from django.conf import settings

logger = logging.getLogger(__name__)

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
    logger.warning("PyPDF2 not installed")

try:
    import docx
except ImportError:
    docx = None
    logger.warning("python-docx not installed")

try:
    from PIL import Image
except ImportError:
    Image = None
    logger.warning("Pillow not installed")


class FileContentAnalyzer:
    """تحلیل محتوای فایل‌ها"""
    
    @staticmethod
    def extract_text_from_pdf(file_path):
        if PyPDF2 is None:
            return None
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages[:10]:
                    text += page.extract_text() or ""
            return text[:3000]
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            return None
    
    @staticmethod
    def extract_text_from_docx(file_path):
        if docx is None:
            return None
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs[:50]])
            return text[:2000]
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            return None
    
    @staticmethod
    def extract_text_from_txt(file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return text[:2000]
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                return text[:2000]
            except Exception as e:
                logger.error(f"Error reading TXT: {e}")
                return None
        except Exception as e:
            logger.error(f"Error reading TXT: {e}")
            return None
    
    @staticmethod
    def extract_metadata_from_image(file_path):
        if Image is None:
            return None
        try:
            with Image.open(file_path) as img:
                metadata = {
                    "format": img.format,
                    "size": img.size,
                    "mode": img.mode,
                    "width": img.width,
                    "height": img.height
                }
                return metadata
        except Exception as e:
            logger.error(f"Error reading image: {e}")
            return None
    
    @staticmethod
    def get_file_summary(file_obj):
        file_path = file_obj.path
        file_extension = os.path.splitext(file_obj.name)[1].lower()
        
        result = {
            "file_name": file_obj.name,
            "file_size": file_obj.size,
            "file_type": file_extension,
            "content_summary": None,
            "key_topics": [],
            "metadata": {}
        }
        
        if file_extension == '.pdf':
            content = FileContentAnalyzer.extract_text_from_pdf(file_path)
            if content:
                result["content_summary"] = content[:500]
                
        elif file_extension == '.docx':
            content = FileContentAnalyzer.extract_text_from_docx(file_path)
            if content:
                result["content_summary"] = content[:500]
                
        elif file_extension == '.txt':
            content = FileContentAnalyzer.extract_text_from_txt(file_path)
            if content:
                result["content_summary"] = content[:500]
                
        elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
            metadata = FileContentAnalyzer.extract_metadata_from_image(file_path)
            if metadata:
                result["metadata"] = metadata
                result["content_summary"] = f"تصویر با فرمت {metadata.get('format')}"
        
        return result


class UserFilesAnalyzer:
    
    def __init__(self, user):
        self.user = user
        self.file_analyzer = FileContentAnalyzer()
    
    def analyze_all_files(self):
        try:
            from apps.posts.models import Post, PostMedia
            
            posts_with_single_file = Post.objects.filter(
                user=self.user,
                file__isnull=False
            ).exclude(file='')
            
            multiple_media = PostMedia.objects.filter(
                post__user=self.user
            ).select_related('post')
            
            files_analysis = []
            all_content = []
            file_types = {}
            total_size = 0
            
            for post in posts_with_single_file:
                try:
                    if post.file and hasattr(post.file, 'path') and os.path.exists(post.file.path):
                        file_info = self.file_analyzer.get_file_summary(post.file)
                        
                        files_analysis.append({
                            "post_id": str(post.id),
                            "post_content": post.content[:100] if post.content else "",
                            "file_info": file_info,
                            "file_source": "single_file",
                            "created_at": post.created_at.isoformat() if post.created_at else None
                        })
                        
                        ext = file_info["file_type"]
                        file_types[ext] = file_types.get(ext, 0) + 1
                        total_size += file_info["file_size"]
                        
                        if file_info.get("content_summary"):
                            all_content.append(file_info["content_summary"])
                            
                except Exception as e:
                    logger.error(f"Error analyzing single file for post {post.id}: {e}")
            
            for media in multiple_media:
                try:
                    if media.file and hasattr(media.file, 'path') and os.path.exists(media.file.path):
                        file_info = self.file_analyzer.get_file_summary(media.file)
                        
                        post = media.post
                        
                        files_analysis.append({
                            "post_id": str(post.id),
                            "post_content": post.content[:100] if post.content else "",
                            "media_id": str(media.id),
                            "file_info": file_info,
                            "file_source": "post_media",
                            "order": media.order,
                            "created_at": post.created_at.isoformat() if post.created_at else None
                        })
                        
                        ext = file_info["file_type"]
                        file_types[ext] = file_types.get(ext, 0) + 1
                        total_size += file_info["file_size"]
                        
                        if file_info.get("content_summary"):
                            all_content.append(file_info["content_summary"])
                            
                except Exception as e:
                    logger.error(f"Error analyzing media file {media.id}: {e}")
            
            total_files = posts_with_single_file.count() + multiple_media.count()
            
            return {
                "total_files": total_files,
                "total_size_bytes": total_size,
                "file_types": file_types,
                "files_analysis": files_analysis[:20],
                "combined_content": " ".join(all_content)[:3000]
            }
            
        except Exception as e:
            logger.error(f"Error in analyze_all_files: {e}")
            return {
                "total_files": 0,
                "total_size_bytes": 0,
                "file_types": {},
                "files_analysis": [],
                "combined_content": "",
                "error": str(e)
            }